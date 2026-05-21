"""PR-378 coverage for CJK font fallback in docx exports.

Field bug 2026-05-13: user opened the manuscript.docx in Microsoft
Word and every Chinese character rendered as ``□`` (tofu / replacement
box). WPS rendered fine because it has a built-in CJK fallback.
Cause: python-docx's ``Document()`` only sets a Latin font (Calibri);
Word looks up CJK characters via ``w:rFonts/w:eastAsia`` which had no
value, so the Latin Calibri's empty glyph map was used.

Fix: inject ``<w:rFonts w:eastAsia="SimSun"/>`` into the document's
default ``rPrDefault`` (styles.xml) so every run inherits a working
CJK font.
"""

from __future__ import annotations

import zipfile
from pathlib import Path


def _read_zip(docx: Path, member: str) -> str:
    with zipfile.ZipFile(docx) as zf, zf.open(member) as f:
        return f.read().decode("utf-8")


def test_docx_styles_xml_has_east_asia_font(tmp_path: Path) -> None:
    """Bug repro: open the produced docx, inspect styles.xml, confirm
    the East Asian default font tag is present."""
    from autoessay.agents.exporter import _write_docx

    out = tmp_path / "manuscript.docx"
    _write_docx(out, "# 标题\n\n中文段落。\n")
    styles = _read_zip(out, "word/styles.xml")
    # PR-378: docDefaults > rPrDefault > rPr > rFonts > w:eastAsia
    assert "w:eastAsia" in styles, (
        "no East Asian font registered in styles.xml; Word will render CJK as □"
    )
    # Must reference an actual CJK font name. ``SimSun`` (宋体) is the
    # universal Windows/Word fallback.
    assert "SimSun" in styles, "expected SimSun as the default East Asian font"


def test_docx_with_cjk_table_cells_renders_each_cell(tmp_path: Path) -> None:
    """Smoke regression: PR-375 markdown-table parsing still works
    AND now the cell text + the 表 N caption ship with a CJK font."""
    from autoessay.agents.exporter import _write_docx

    md = (
        "# 测试\n"
        "\n"
        "| 变量 | 含义 |\n"
        "|---|---|\n"
        "| 教育实践_i | 是否参与教学、修辞训练或课程安排 |\n"
        "| 文本生产_i | 是否参与写作、编辑、翻译或出版 |\n"
    )
    out = tmp_path / "with_cjk_table.docx"
    _write_docx(out, md)
    body = _read_zip(out, "word/document.xml")
    styles = _read_zip(out, "word/styles.xml")
    # Table is real.
    assert "<w:tbl>" in body
    # Caption present.
    assert "表 1" in body
    # CJK cell content present (NOT tofu).
    for needle in (
        "变量",
        "含义",
        "教育实践_i",
        "是否参与教学",
        "文本生产_i",
        "是否参与写作",
    ):
        assert needle in body, f"CJK cell text missing: {needle}"
    # East Asian font injection survived.
    assert 'w:eastAsia="SimSun"' in styles


def test_inject_cjk_default_font_is_idempotent(tmp_path: Path) -> None:
    """Calling the injector twice should not duplicate the
    ``w:rFonts`` element. Future callers (e.g. literature_usage docx)
    might double-call by accident."""
    from docx import Document

    from autoessay.agents.exporter import _inject_cjk_default_font

    doc = Document()
    _inject_cjk_default_font(doc)
    _inject_cjk_default_font(doc)  # second call should no-op
    out = tmp_path / "double_call.docx"
    doc.save(str(out))
    styles = _read_zip(out, "word/styles.xml")
    # Count ``w:rFonts`` elements inside docDefaults; should appear
    # at most once for the default + at most once for Normal style.
    # Idempotency means we don't keep appending new elements on each
    # call.
    east_asia_count = styles.count('w:eastAsia="SimSun"')
    assert east_asia_count in {1, 2}, (
        f"expected at most 2 SimSun east-asia entries (docDefaults + Normal style); "
        f"got {east_asia_count}"
    )


async def test_export_download_sets_no_store_cache_header(app_session) -> None:  # type: ignore[no-untyped-def]
    """PR-378: defeat the Cloudflare 4h cache that was serving users
    stale docx after backend fixes deployed. Live-discovered when
    ``cf-cache-status: HIT`` + ``cache-control: max-age=14400`` kept
    feeding the old pipe-paragraph docx to ``run_e11d7e52``'s owner
    even after PR-375 wrote a new file to disk."""
    from datetime import datetime, timezone

    from httpx import ASGITransport, AsyncClient
    from sqlalchemy import select

    from autoessay.main import app
    from autoessay.models import Run

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        project_response = await client.post(
            "/api/projects",
            json={
                "title": "Cache header test",
                "domain_id": "financial_history",
                "target_journal": None,
            },
        )
        assert project_response.status_code == 201
        project_id = project_response.json()["id"]
        created = await client.post(f"/api/projects/{project_id}/runs")
        assert created.status_code == 201
        run_id = created.json()["id"]
        # Materialize a fake export on disk.
        with app_session() as session:
            run = session.scalar(select(Run).where(Run.id == run_id))
            assert run is not None
            run.updated_at = datetime.now(timezone.utc)
            exports_dir = Path(run.run_dir) / "exports"
            exports_dir.mkdir(parents=True, exist_ok=True)
            (exports_dir / "manuscript.docx").write_bytes(b"PK\x03\x04fake-docx")
            session.commit()
        response = await client.get(
            f"/api/runs/{run_id}/exports/manuscript.docx",
        )
        assert response.status_code == 200
        # No-store family of directives so Cloudflare + browsers
        # never cache the binary.
        cache_control = response.headers.get("cache-control", "")
        assert "no-store" in cache_control
        # Some Cloudflare configurations honour max-age over no-store
        # if max-age isn't zero; pin max-age=0 too.
        assert "max-age=0" in cache_control
        # Belt-and-braces legacy headers for ancient proxies.
        assert response.headers.get("pragma", "").lower() == "no-cache"


def test_inject_accepts_custom_font_name(tmp_path: Path) -> None:
    """Sanity: caller can override the font name via kwarg if a deploy
    target needs Noto Sans CJK SC or Microsoft YaHei instead."""
    from docx import Document

    from autoessay.agents.exporter import _inject_cjk_default_font

    doc = Document()
    _inject_cjk_default_font(doc, east_asia_font="Noto Sans CJK SC")
    out = tmp_path / "custom_font.docx"
    doc.save(str(out))
    styles = _read_zip(out, "word/styles.xml")
    assert "Noto Sans CJK SC" in styles
