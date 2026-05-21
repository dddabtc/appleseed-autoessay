"""Final artifact exporter."""

from __future__ import annotations

import hashlib
import json
import re
import subprocess
from collections.abc import Mapping, Sequence
from datetime import datetime, timezone
from pathlib import Path

import markdown
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import ValidationError
from sqlalchemy import select
from sqlalchemy.orm import Session

from autoessay.agents.critic import latest_draft_dir, run_citation_audit
from autoessay.agents.final_rewrite import (
    attempt_exports_policy_polish_retry,
    load_latest_rewrite_artifact,
)
from autoessay.agents.literature_usage import build_literature_usage_table
from autoessay.agents.manuscript_compose import compose_manuscript, strip_existing_paper_matter
from autoessay.agents.self_check import (
    render_self_check_markdown,
    report_to_dict,
    run_self_check,
)
from autoessay.clients.common import NormalizedSource
from autoessay.config import get_settings
from autoessay.db import SessionLocal
from autoessay.models import Checkpoint, Project, Run, User
from autoessay.state_machine import InvalidTransition, append_event, assert_run_active, transition

EXPORT_FORMATS = {"markdown", "docx", "html", "latex", "bibtex", "csl_json"}
DEFAULT_EXPORT_FORMATS = ["markdown", "docx", "html", "latex", "bibtex", "csl_json"]


def run_exports(
    run_id: str,
    db_session: Session | None = None,
    *,
    lock_token: str | None = None,
) -> dict[str, object]:
    """Run exports. Stage 3.E follow-up P0: ``lock_token`` triggers
    owner-checked phase-start lock release at exit. PR-A4.1b
    (2026-05-02): wraps in ``maybe_run_with_versioning``."""
    from autoessay.phase_lock import phase_lock_release_on_exit
    from autoessay.phase_version import maybe_run_with_versioning

    def _execute(session: Session) -> dict[str, object]:
        run = session.scalar(select(Run).where(Run.id == run_id))
        if run is None:
            raise ValueError(f"run not found: {run_id}")
        result: dict[str, object] = {}

        def _runner() -> None:
            result["value"] = _run_exports_with_session(run_id, session)

        maybe_run_with_versioning(session, run, "exports", _runner)
        return result.get("value", {})  # type: ignore[return-value]

    with phase_lock_release_on_exit(run_id, "exports", lock_token, session=db_session):
        if db_session is not None:
            return _execute(db_session)
        with SessionLocal() as session:
            return _execute(session)


def load_exports_payload(run: Run) -> dict[str, object]:
    exports_dir = Path(run.run_dir) / "exports"
    manifest = _load_json_mapping(exports_dir / "manifest.json")
    return {
        "run_id": run.id,
        "manifest": manifest,
        "files": _download_links(run.id, manifest),
    }


def _run_exports_with_session(run_id: str, session: Session) -> dict[str, object]:
    run = session.scalar(select(Run).where(Run.id == run_id))
    if run is None:
        raise ValueError(f"run not found: {run_id}")
    assert_run_active(run, session)
    if run.state != "USER_FINAL_ACCEPTANCE":
        raise InvalidTransition(f"Exporter requires USER_FINAL_ACCEPTANCE, got {run.state}")
    project = session.scalar(select(Project).where(Project.id == run.project_id))
    if project is None:
        return _fail_fixable(run, session, "Exporter could not find the project.")
    # PR-G-Regressions A+B: replace the legacy ``project.language``
    # direct-read with ``_resolve_paper_language(project, kernel)`` —
    # the same auto-detection drafter has used since PR-256. Without
    # this, NewRunPage default ``en`` + Chinese kernel produced
    # mixed-language manuscripts (drafter wrote zh body, exporter
    # generated en title/abstract/keywords prepended on top).
    from autoessay.agents.drafter import _resolve_paper_language

    raw_kernel = getattr(run, "research_kernel_json", None)
    if isinstance(raw_kernel, str):
        import json as _json

        try:
            raw_kernel = _json.loads(raw_kernel)
        except (ValueError, TypeError):
            raw_kernel = None
    project_language = _resolve_paper_language(project, raw_kernel) if project else "en"

    transition(run, "EXPORTS_RUNNING", session, reason="Exports started")
    append_event(session, run, "phase_started", {"phase": "exports", "run_id": run.id})
    session.commit()
    session.refresh(run)

    run_dir = Path(run.run_dir)
    policy_attempts: list[dict[str, object]] = []
    max_policy_retries = int(getattr(get_settings(), "exports_policy_max_polish_retries", 2))
    for retry_index in range(max_policy_retries + 1):
        blocking_issues = _load_blocking_issues(run_dir / "reviews" / "blocking_issues.json")
        audit_rows, audit_blockers = run_citation_audit(run_dir)
        # Stage 3.E follow-up: force-approve marked some BLOCKERs as
        # ``resolved_by="user_force_approve"`` in ``blocking_issues.json``.
        # The fresh ``run_citation_audit`` re-scans claim_map.jsonl and
        # would naively re-create the same blockers — which sends the
        # run back to FAILED_POLICY in a loop. Drop audit_blockers whose
        # paragraph_id was already user-resolved so force-approve is
        # actually honored across re-export attempts.
        user_resolved_paragraphs = {
            str(issue.get("paragraph_id"))
            for issue in blocking_issues
            if issue.get("resolved")
            and issue.get("paragraph_id")
            and issue.get("resolved_by") == "user_force_approve"
        }
        if user_resolved_paragraphs:
            audit_blockers = [
                blocker
                for blocker in audit_blockers
                if str(getattr(blocker, "paragraph_id", "") or "") not in user_resolved_paragraphs
            ]
        unresolved = [issue for issue in blocking_issues if not bool(issue.get("resolved"))]
        if not unresolved and not audit_blockers:
            break
        guidance = _policy_guidance(unresolved, audit_blockers)
        if retry_index >= max_policy_retries:
            return _fail_policy(run, session, guidance, audit_rows, policy_attempts)
        repair_result = attempt_exports_policy_polish_retry(
            run=run,
            project=project,
            session=session,
            guidance=guidance,
            failure_class="failed_policy",
            retry_index=retry_index + 1,
            audit_rows=audit_rows,
        )
        policy_attempts.append(
            {
                "retry_index": retry_index + 1,
                "guidance": guidance,
                "unresolved_count": len(unresolved),
                "audit_blocker_count": len(audit_blockers),
                "repair_result": repair_result,
            },
        )
        if repair_result.get("status") != "rewritten":
            return _fail_policy(run, session, guidance, audit_rows, policy_attempts)
        if unresolved:
            _mark_blocking_issues_auto_polish_attempted(
                run_dir / "reviews" / "blocking_issues.json",
                unresolved,
                retry_index=retry_index + 1,
                guidance=guidance,
            )

    draft_dir = latest_draft_dir(run_dir)
    if draft_dir is None:
        return _fail_fixable(run, session, "Exporter could not find a draft directory.")
    body_markdown, claim_map = _final_manuscript_and_claim_map(run_dir, draft_dir)
    if not body_markdown.strip():
        return _fail_fixable(run, session, "Exporter could not find a final manuscript.")
    citations_bib = _read_optional_text(draft_dir / "citations.bib")
    shortlist = _read_sources_json(run_dir / "sources" / "shortlist.json")
    formats = _export_formats(session, run)
    cited_source_ids = _cited_source_ids(claim_map)
    cited_sources = [s for s in shortlist if s.source_id in cited_source_ids]
    body_markdown = strip_existing_paper_matter(body_markdown, project_language)
    # PR-G-Regressions C2 (codex round-1 amendment Q1): second-pass
    # citation normalize on the final manuscript. The drafter-side
    # ``_normalize_inline_citations_zh`` already runs once, but
    # stylist (and future polish) can re-introduce raw cite markers
    # like ``[crossref:DOI]`` or ``[https://openalex.org/W…]``.
    # Running normalize a second time at the export boundary
    # guarantees the final artifact ships with ``[N]`` markers only.
    if project_language in ("zh", "ja"):
        from autoessay.agents.drafter import _normalize_inline_citations_zh

        body_markdown = _normalize_inline_citations_zh(body_markdown, cited_sources)
        body_markdown = _repair_numeric_citations_from_claim_map(
            body_markdown,
            claim_map=claim_map,
            cited_sources=cited_sources,
        )
    selected_thesis = _load_json_mapping(run_dir / "novelty" / "selected_thesis.json")
    authors = _resolve_authors(session, project)
    manuscript = compose_manuscript(
        run=run,
        session=session,
        body_markdown=body_markdown,
        project_title=project.title if project else "",
        project_language=project_language,
        authors=authors,
        cited_sources=cited_sources,
        selected_thesis=selected_thesis or None,
    )

    exports_dir = run_dir / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    artifacts: dict[str, str] = {}
    if "markdown" in formats:
        _write_text(exports_dir / "manuscript.md", manuscript)
        artifacts["markdown"] = "exports/manuscript.md"
    if "docx" in formats:
        _write_docx(exports_dir / "manuscript.docx", manuscript)
        artifacts["docx"] = "exports/manuscript.docx"
    if "html" in formats:
        _write_html(exports_dir / "manuscript.html", manuscript, project_language)
        artifacts["html"] = "exports/manuscript.html"
    if "latex" in formats:
        _write_latex(exports_dir / "manuscript.tex", manuscript, project_language)
        artifacts["latex"] = "exports/manuscript.tex"
    if "bibtex" in formats:
        _write_text(exports_dir / "citations.bib", citations_bib)
        artifacts["bibtex"] = "exports/citations.bib"
    if "csl_json" in formats:
        _write_json(
            exports_dir / "citations.csl.json",
            _csl_items(shortlist, _cited_source_ids(claim_map), project_language),
        )
        artifacts["csl_json"] = "exports/citations.csl.json"

    # Literature usage table — paper-quality-spec.md §五.11. Always
    # written when there is at least one cited source, regardless of
    # which export formats the user ticked.
    source_notes = _read_synthesizer_source_notes(run_dir)
    usage_table = build_literature_usage_table(
        cited_sources=cited_sources,
        claim_map=claim_map,
        source_notes=source_notes,
        project_language=project_language,
    )
    if usage_table.strip():
        usage_path = exports_dir / "literature_usage_table.md"
        _write_text(usage_path, usage_table)
        artifacts["literature_usage_table"] = "exports/literature_usage_table.md"

    # Self-check report — paper-quality-spec.md §七, 13 items. Always
    # written. Stub mode (AUTOESSAY_SELF_CHECK_STUB=1) returns a
    # report with all items "incomplete" without an LLM call.
    report = run_self_check(
        run=run,
        session=session,
        manuscript_markdown=manuscript,
        project_language=project_language,
    )
    self_check_md = render_self_check_markdown(report, project_language)
    self_check_path = exports_dir / "self_check_report.md"
    _write_text(self_check_path, self_check_md)
    artifacts["self_check_report"] = "exports/self_check_report.md"
    _write_json(exports_dir / "self_check_report.json", report_to_dict(report))
    artifacts["self_check_report_json"] = "exports/self_check_report.json"

    manifest = _manifest_payload(run_dir, artifacts, project_language)
    if policy_attempts:
        manifest["exports_policy_polish_retries"] = policy_attempts
    _write_json(exports_dir / "manifest.json", manifest)
    transition(
        run,
        "EXPORTS_DONE",
        session,
        reason="Exports completed",
        payload={"exports": artifacts},
    )
    append_event(session, run, "phase_done", {"phase": "exports", "exports": artifacts})
    session.commit()
    return {"run_id": run.id, "state": run.state, "exports": artifacts, "manifest": manifest}


def _fail_policy(
    run: Run,
    session: Session,
    guidance: str,
    audit_rows: Sequence[Mapping[str, object]],
    policy_attempts: Sequence[Mapping[str, object]] = (),
) -> dict[str, object]:
    transition(
        run,
        "FAILED_POLICY",
        session,
        reason="Exporter citation gate blocked",
        payload={"guidance": guidance, "exports_policy_polish_retries": list(policy_attempts)},
    )
    append_event(
        session,
        run,
        "phase_failed",
        {
            "phase": "exports",
            "failure_class": "failed_policy",
            "guidance": guidance,
            "claim_audit": list(audit_rows),
            "exports_policy_polish_retries": list(policy_attempts),
        },
    )
    session.commit()
    return {"run_id": run.id, "state": run.state, "guidance": guidance}


def _fail_fixable(run: Run, session: Session, guidance: str) -> dict[str, object]:
    transition(
        run,
        "FAILED_FIXABLE",
        session,
        reason="Exporter missing fixable artifact",
        payload={"guidance": guidance},
    )
    append_event(
        session,
        run,
        "phase_failed",
        {"phase": "exports", "failure_class": "failed_fixable", "guidance": guidance},
    )
    session.commit()
    return {"run_id": run.id, "state": run.state, "guidance": guidance}


def _load_blocking_issues(path: Path) -> list[dict[str, object]]:
    payload = _load_json_mapping(path)
    raw_issues: object = payload.get("issues")
    if not isinstance(raw_issues, list):
        return []
    return [issue for issue in raw_issues if isinstance(issue, dict)]


def _mark_blocking_issues_auto_polish_attempted(
    path: Path,
    unresolved: Sequence[Mapping[str, object]],
    *,
    retry_index: int,
    guidance: str,
) -> None:
    payload = _load_json_mapping(path)
    raw_issues: object = payload.get("issues")
    if not isinstance(raw_issues, list):
        return
    unresolved_ids = {
        str(issue.get("issue_id") or issue.get("id") or issue.get("description") or "")
        for issue in unresolved
    }
    changed = False
    for issue in raw_issues:
        if not isinstance(issue, dict) or issue.get("resolved"):
            continue
        issue_key = str(issue.get("issue_id") or issue.get("id") or issue.get("description") or "")
        if issue_key not in unresolved_ids:
            continue
        issue["resolved"] = True
        issue["resolved_by"] = "auto_polish_retry"
        issue["resolved_at"] = datetime.now(timezone.utc).isoformat()
        issue["resolution_note"] = guidance
        issue["auto_polish_retry_index"] = retry_index
        changed = True
    if changed:
        payload["issues"] = raw_issues
        _write_json(path, payload)


def _policy_guidance(
    unresolved: Sequence[Mapping[str, object]],
    audit_blockers: Sequence[object],
) -> str:
    if audit_blockers:
        first = audit_blockers[0]
        description = getattr(first, "description", "A citation has no DOI or URL.")
        return (
            f"{description} Add a DOI or URL, upload the source, or mark the source "
            "as unverified-user-supplied before exporting."
        )
    if unresolved:
        first = unresolved[0]
        description = first.get("description")
        if isinstance(description, str) and description:
            return description
    return "Unresolved blocking citation issues remain before export."


def _final_manuscript_and_claim_map(
    run_dir: Path,
    draft_dir: Path,
) -> tuple[str, list[dict[str, object]]]:
    rewrite = load_latest_rewrite_artifact(run_dir)
    if rewrite is not None:
        return rewrite.manuscript, [dict(item) for item in rewrite.claim_map]
    return _final_manuscript(draft_dir), _load_jsonl_objects(draft_dir / "claim_map.jsonl")


def _final_manuscript(draft_dir: Path) -> str:
    styled = _read_optional_text(draft_dir / "style" / "paper_styled.md")
    if styled.strip():
        return styled
    return _read_optional_text(draft_dir / "manuscript.md")


def _repair_numeric_citations_from_claim_map(
    body_markdown: str,
    *,
    claim_map: Sequence[Mapping[str, object]],
    cited_sources: Sequence[NormalizedSource],
) -> str:
    """Align final numeric markers with paragraph-level claim_map sources.

    LLM stages sometimes emit numeric markers based on shortlist
    positions (for example ``[6][7]``) even though the final reference
    list is numbered from the cited-source subset. At export, claim_map
    is the authoritative mapping from paragraphs to source_ids, so use
    it to make the body markers and reference list one-to-one before
    composing front/back matter.
    """
    if not body_markdown.strip() or not cited_sources:
        return body_markdown
    source_to_tag = {
        source.source_id: f"[{index}]" for index, source in enumerate(cited_sources, 1)
    }
    paragraph_tags = _claim_map_paragraph_tags(claim_map, source_to_tag)
    if not paragraph_tags:
        return body_markdown

    paragraphs = body_markdown.split("\n\n")
    repaired: list[str] = []
    claim_index = 0
    max_ref = len(cited_sources)
    for paragraph in paragraphs:
        stripped = paragraph.strip()
        if not stripped:
            repaired.append(paragraph)
            continue
        if _is_markdown_heading(stripped) or stripped == "---":
            repaired.append(paragraph)
            continue
        expected_tags = paragraph_tags[claim_index] if claim_index < len(paragraph_tags) else []
        claim_index += 1
        paragraph = _replace_raw_source_id_markers(paragraph, source_to_tag)
        current_nums = [int(value) for value in re.findall(r"\[(\d{1,3})\]", paragraph)]
        if expected_tags:
            if current_nums:
                repaired.append(_replace_numeric_citation_clusters(paragraph, expected_tags))
            else:
                repaired.append(_append_tags_to_paragraph(paragraph, expected_tags))
        elif any(number > max_ref for number in current_nums):
            repaired.append(re.sub(r"(?:\s*\[\d{1,3}\])+", "", paragraph).rstrip())
        else:
            repaired.append(paragraph)
    return "\n\n".join(repaired).rstrip() + "\n"


def _claim_map_paragraph_tags(
    claim_map: Sequence[Mapping[str, object]],
    source_to_tag: Mapping[str, str],
) -> list[list[str]]:
    grouped: dict[str, list[str]] = {}
    order: list[str] = []
    for index, claim in enumerate(claim_map, 1):
        paragraph_id = str(claim.get("paragraph_id") or f"claim-{index:03d}")
        if paragraph_id not in grouped:
            grouped[paragraph_id] = []
            order.append(paragraph_id)
        raw_source_ids = claim.get("source_ids")
        if not isinstance(raw_source_ids, list):
            continue
        for source_id in raw_source_ids:
            if not isinstance(source_id, str):
                continue
            tag = source_to_tag.get(source_id)
            if tag and tag not in grouped[paragraph_id]:
                grouped[paragraph_id].append(tag)
    return [grouped[paragraph_id] for paragraph_id in order]


def _replace_raw_source_id_markers(paragraph: str, source_to_tag: Mapping[str, str]) -> str:
    repaired = paragraph
    bracket_patterns = (
        r"\[\s*([^\]]+?)\s*\]",
        r"［\s*([^］]+?)\s*］",
        r"【\s*([^】]+?)\s*】",
        r"〔\s*([^〕]+?)\s*〕",
        r"[（(]\s*([^）)]+?)\s*[）)]",
    )
    for pattern in bracket_patterns:
        repaired = re.sub(
            pattern,
            lambda match: _replace_composite_source_id_marker(match, source_to_tag),
            repaired,
        )
    ordered_sources = sorted(source_to_tag.items(), key=lambda item: len(item[0]), reverse=True)
    for source_id, tag in ordered_sources:
        escaped = re.escape(source_id)
        repaired = re.sub(rf"[\[［【]\s*{escaped}\s*[\]］】]", tag, repaired)
        repaired = re.sub(rf"[（(]\s*{escaped}\s*[）)]", tag, repaired)
    return repaired


def _replace_composite_source_id_marker(
    match: re.Match[str],
    source_to_tag: Mapping[str, str],
) -> str:
    inner = match.group(1).strip()
    pieces = _split_source_marker_pieces(inner)
    if len(pieces) < 2:
        return match.group(0)
    tags: list[str] = []
    for piece in pieces:
        tag = source_to_tag.get(piece)
        if tag is None:
            return match.group(0)
        tags.append(tag)
    return "".join(tags)


def _split_source_marker_pieces(inner: str) -> list[str]:
    return [
        piece.strip() for piece in re.split(r"\s*(?:[;；,，、]|\s+)\s*", inner) if piece.strip()
    ]


def _replace_numeric_citation_clusters(paragraph: str, expected_tags: Sequence[str]) -> str:
    canonical = "".join(expected_tags)
    replaced = False

    def repl(_match: re.Match[str]) -> str:
        nonlocal replaced
        if not replaced:
            replaced = True
            return canonical
        return ""

    return re.sub(r"(?:\[\d{1,3}\])+", repl, paragraph)


def _append_tags_to_paragraph(paragraph: str, expected_tags: Sequence[str]) -> str:
    canonical = "".join(expected_tags)
    stripped = paragraph.rstrip()
    if not stripped:
        return paragraph
    if stripped[-1] in "。！？.!?":
        return stripped[:-1] + canonical + stripped[-1]
    return stripped + canonical


def _is_markdown_heading(text: str) -> bool:
    return bool(re.match(r"^\s*#{1,6}\s+", text))


def _export_formats(session: Session, run: Run) -> list[str]:
    checkpoint = session.scalar(
        select(Checkpoint)
        .where(Checkpoint.run_id == run.id)
        .where(Checkpoint.checkpoint_type == "USER_FINAL_ACCEPTANCE")
        .order_by(Checkpoint.created_at.desc(), Checkpoint.id.desc())
        .limit(1),
    )
    if checkpoint is None:
        return list(DEFAULT_EXPORT_FORMATS)
    payload = _json_object(checkpoint.decision_payload)
    raw_formats = payload.get("export_formats")
    if not isinstance(raw_formats, list):
        return list(DEFAULT_EXPORT_FORMATS)
    formats = [item for item in raw_formats if isinstance(item, str) and item in EXPORT_FORMATS]
    return formats or list(DEFAULT_EXPORT_FORMATS)


_MD_TABLE_SEPARATOR_RE = re.compile(r"^\|[\s:|-]+\|$")
_MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
# PR-381: tolerate trailing citation markers after the closing pipe,
# e.g. ``| ... | ... |[2]`` or ``| ... | ... | [12][34]``. LLM emits
# row-scoped citations after the terminal pipe; previous parser
# rejected those rows because ``endsWith("|")`` was False, leaving
# the last row of every cited table as garbled pipe-text.
# Codex AGREE-WITH-AMENDMENTS PR-381 #2: pin the trailing-citation
# regex to numeric-only ``[N]`` / ``[12]`` / ``[3][4]`` groups so
# table cells whose RIGHTMOST cell happens to end on a bracket
# (e.g. ``[evidence: §3.2]``) don't get their bracket peeled off
# by mistake. LLM-emitted row-scoped citations are always numeric.
_MD_TABLE_TRAILING_CITATION_RE = re.compile(r"(\s*(?:\[\s*\d+\s*\])+\s*)$")


def _strip_trailing_table_citations(line: str) -> str:
    """Strip ``[N]``/``[12]``/``[3][4]`` groups from the end of a
    table row when they appear AFTER the closing pipe. Returns the
    line unchanged if the trailing brackets are inside a normal cell."""
    s = line
    while True:
        m = _MD_TABLE_TRAILING_CITATION_RE.search(s)
        if m is None:
            break
        # Only strip if what precedes the brackets is a closing pipe
        # (possibly with whitespace); otherwise the brackets are
        # part of cell content.
        preceding = s[: m.start()].rstrip()
        if not preceding.endswith("|"):
            break
        s = preceding
    return s


def _is_table_row(line: str) -> bool:
    """A table row starts with ``|``, ends with ``|`` after stripping
    trailing ``[N]`` citations, and has at least one inner ``|``."""
    if not line.startswith("|"):
        return False
    stripped = _strip_trailing_table_citations(line)
    return stripped.startswith("|") and stripped.endswith("|") and len(stripped) >= 2


def _split_table_cells(line: str) -> list[str]:
    """Split a (possibly citation-trailing) table row into cells."""
    stripped = _strip_trailing_table_citations(line)
    return [c.strip() for c in stripped.strip("|").split("|")]


def _looks_like_table_header(lines: list[str], i: int) -> bool:
    """Return True iff ``lines[i]`` and ``lines[i+1]`` together form a
    valid markdown-table header + separator pair."""
    if i + 1 >= len(lines):
        return False
    head = lines[i].strip()
    if not _is_table_row(head):
        return False
    sep = lines[i + 1].strip()
    return bool(_MD_TABLE_SEPARATOR_RE.match(sep))


def _parse_md_table(lines: list[str], i: int) -> tuple[list[str], list[list[str]], int]:
    """Parse a markdown table starting at ``lines[i]``.

    Returns ``(header_cells, data_rows, lines_consumed)``. Assumes the
    caller already verified ``_looks_like_table_header(lines, i)``.
    """
    header = _split_table_cells(lines[i].strip())
    j = i + 2  # skip separator
    data_rows: list[list[str]] = []
    while j < len(lines):
        sub = lines[j].strip()
        if not _is_table_row(sub):
            break
        data_rows.append(_split_table_cells(sub))
        j += 1
    return header, data_rows, j - i


def _inject_cjk_default_font(document: object, *, east_asia_font: str = "SimSun") -> None:
    """PR-378: register a default East Asian font on the document so
    Microsoft Word renders CJK characters as glyphs instead of empty
    placeholder boxes (``□``).

    python-docx's ``Document()`` only sets a Latin font (Calibri).
    Word looks up CJK characters under ``w:rFonts/w:eastAsia``; with
    no value set, it falls back to the default Latin font, which on
    macOS/Win11 Word doesn't contain CJK glyphs → every Chinese
    character renders as a tofu box. WPS has a built-in CJK fallback
    so it renders fine.

    Fix: inject ``<w:rFonts w:eastAsia="SimSun"/>`` into the
    document's default ``rPrDefault`` (``word/styles.xml``). Every
    run that doesn't explicitly set its own font inherits this.
    SimSun (宋体) is the OpenType Chinese font that ships with every
    Word install on every platform, so this is a portable default.

    Field-discovered 2026-05-13 when user reported tables "even more
    broken" in Word vs WPS — the tables were already PR-375-fixed,
    but Word's CJK fallback was the second bug masking it.
    """
    styles_element = document.styles.element  # type: ignore[attr-defined]
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    # Set ALL four font-region attributes so Word picks the CJK font
    # for Chinese codepoints AND keeps Latin/numeric prose on the
    # default Calibri. ``w:eastAsia`` is the one Word actually checks
    # for CJK lookups; ``w:hAnsi`` and ``w:ascii`` keep the Latin
    # fallback unchanged; ``w:cs`` covers complex-script fallback so
    # Arabic / Thai etc. still hit Word's defaults.
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    # Also pin the Normal style to use the same East Asian font so
    # docx readers that prefer style-defined fonts over docDefaults
    # (older Word versions, some Pages.app builds) get a consistent
    # result.
    try:
        normal_style = document.styles["Normal"]  # type: ignore[attr-defined]
        normal_rpr = normal_style.element.find(qn("w:rPr"))
        if normal_rpr is None:
            normal_rpr = OxmlElement("w:rPr")
            normal_style.element.append(normal_rpr)
        normal_rfonts = normal_rpr.find(qn("w:rFonts"))
        if normal_rfonts is None:
            normal_rfonts = OxmlElement("w:rFonts")
            normal_rpr.append(normal_rfonts)
        normal_rfonts.set(qn("w:eastAsia"), east_asia_font)
    except (KeyError, AttributeError):
        # The Normal style isn't strictly required for our purpose;
        # the docDefaults injection above is the load-bearing fix.
        pass


def _write_docx(path: Path, manuscript: str) -> None:
    """PR-375: parse markdown tables into python-docx ``Table``
    objects, parse ``![alt](url)`` images into placeholder paragraphs
    with auto-numbered captions. Tables get a bold ``表 N`` caption
    ABOVE; figures get a ``图 N`` caption BELOW (codex
    AGREE-WITH-AMENDMENTS 2026-05-13 amendment 1 — Chinese journal
    convention). Numbers are assigned in document order; sidecar /
    manifest unchanged.

    PR-378: inject a CJK default font so Word renders Chinese as
    glyphs instead of ``□``.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _inject_cjk_default_font(document)
    lines = manuscript.splitlines()
    i = 0
    table_counter = 0
    figure_counter = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped or re.match(r'^<a id="[^"]+"></a>$', stripped):
            i += 1
            continue
        # PR-381: ``$$ ... $$`` display math block → centered
        # monospace paragraph. python-docx has no native math
        # typesetting (OMath is non-trivial), but at minimum the
        # ``$$`` markers must be consumed so the formula text
        # doesn't render as raw ``$$`` literals around a prose
        # paragraph. Buffer everything between the open/close
        # ``$$`` and emit a single formula paragraph.
        if stripped == "$$" or (stripped.startswith("$$") and not stripped.endswith("$$")):
            math_lines: list[str] = []
            if stripped != "$$":
                math_lines.append(stripped[2:].strip())
            i += 1
            while i < len(lines):
                sub = lines[i].strip()
                if sub == "$$" or sub.endswith("$$"):
                    if sub != "$$":
                        math_lines.append(re.sub(r"\$\$\s*$", "", sub).strip())
                    i += 1
                    break
                math_lines.append(lines[i])
                i += 1
            formula = "\n".join(s for s in (m.rstrip() for m in math_lines) if s)
            if formula:
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                math_para = document.add_paragraph()
                math_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                math_run = math_para.add_run(formula)
                # Codex AGREE-WITH-AMENDMENTS PR-381 #1: ditch the
                # blanket italic — too aggressive for CJK / mixed
                # content. Cambria Math + center alignment is enough
                # to visually distinguish a formula paragraph.
                math_run.font.name = "Cambria Math"
            continue
        # Inline single-line ``$$...$$``: same treatment.
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            formula = stripped[2:-2].strip()
            if formula:
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                math_para = document.add_paragraph()
                math_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                math_run = math_para.add_run(formula)
                # Codex AGREE-WITH-AMENDMENTS PR-381 #1: ditch the
                # blanket italic — too aggressive for CJK / mixed
                # content. Cambria Math + center alignment is enough
                # to visually distinguish a formula paragraph.
                math_run.font.name = "Cambria Math"
            i += 1
            continue
        # Markdown table → python-docx Table + 表 N caption ABOVE.
        if _looks_like_table_header(lines, i):
            header, data_rows, advance = _parse_md_table(lines, i)
            table_counter += 1
            cap = document.add_paragraph()
            cap_run = cap.add_run(f"表 {table_counter}")
            cap_run.bold = True
            num_cols = max(len(header), 1)
            tbl = document.add_table(rows=1 + len(data_rows), cols=num_cols)
            # Try a built-in border-bearing style; fall back silently
            # so older python-docx without the style still renders a
            # bordered table.
            for style_name in ("Light Grid Accent 1", "Light Grid", "Table Grid"):
                try:
                    tbl.style = style_name
                    break
                except (KeyError, ValueError):
                    continue
            for j, cell in enumerate(header[:num_cols]):
                tbl.rows[0].cells[j].text = cell
                for run in tbl.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True
            for r, row in enumerate(data_rows, start=1):
                padded = row + [""] * (num_cols - len(row))
                for j, cell in enumerate(padded[:num_cols]):
                    tbl.rows[r].cells[j].text = cell
            i += advance
            continue
        # Markdown image ``![alt](url)`` → placeholder paragraph +
        # ``图 N`` caption BELOW (codex amendment 1: figures get
        # caption below per Chinese convention).
        img_match = _MD_IMAGE_RE.fullmatch(stripped)
        if img_match:
            figure_counter += 1
            alt, _url = img_match.group(1), img_match.group(2)
            placeholder = document.add_paragraph()
            placeholder_run = placeholder.add_run(
                f"[图 {figure_counter}: {alt or '未命名插图'}]"
                if alt
                else f"[图 {figure_counter}]",
            )
            placeholder_run.italic = True
            cap = document.add_paragraph()
            cap_run = cap.add_run(f"图 {figure_counter}{('  ' + alt) if alt else ''}")
            cap_run.bold = True
            i += 1
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)
        i += 1
    document.save(str(path))


def _add_captions_to_html(html_body: str) -> str:
    """PR-375: post-process markdown→HTML output to wrap ``<table>``
    elements with ``<figure class="table-figure">`` plus a
    ``<figcaption>表 N</figcaption>`` ABOVE (Chinese convention,
    codex AGREE-WITH-AMENDMENTS amendment 1), and ``<img>`` elements
    with ``<figure class="image-figure">`` + ``<figcaption>图 N</figcaption>``
    BELOW. Numbering is sequential in document order across both
    types (separate counters per type).
    """
    table_counter = 0
    figure_counter = 0

    def replace_table(match: re.Match[str]) -> str:
        nonlocal table_counter
        table_counter += 1
        return (
            f'<figure class="table-figure">'
            f'<figcaption class="table-caption"><strong>表 {table_counter}</strong></figcaption>'
            f"{match.group(0)}"
            f"</figure>"
        )

    def replace_image(match: re.Match[str]) -> str:
        nonlocal figure_counter
        figure_counter += 1
        # Pull alt for the figcaption if present.
        alt_match = re.search(r'alt="([^"]*)"', match.group(0))
        alt = alt_match.group(1) if alt_match else ""
        cap_text = f"<strong>图 {figure_counter}</strong>"
        if alt:
            cap_text += f"  {alt}"
        return (
            f'<figure class="image-figure">'
            f"{match.group(0)}"
            f'<figcaption class="image-caption">{cap_text}</figcaption>'
            f"</figure>"
        )

    # Tables: match the full ``<table>...</table>`` block (greedy +
    # DOTALL so nested rows are captured).
    out = re.sub(r"<table[\s\S]*?</table>", replace_table, html_body)
    # Standalone images (not already inside a <figure>). Match
    # ``<p><img.../></p>`` if python-markdown wrapped it, else bare
    # ``<img/>``. Skip image tags already wrapped by our table pass
    # (won't happen since tables don't contain images here).
    out = re.sub(
        r"<p>\s*(<img[^>]*?/?>)\s*</p>",
        lambda m: replace_image(m),
        out,
    )
    return out


def _write_html(path: Path, manuscript: str, language: str = "en") -> None:
    body = markdown.markdown(
        manuscript,
        extensions=["extra", "sane_lists", "smarty"],
        output_format="html5",
    )
    body = _add_captions_to_html(body)
    lang_attr = (language or "en").strip().lower() or "en"
    html = (
        "<!doctype html>\n"
        f'<html lang="{lang_attr}">\n'
        "<head>\n"
        '  <meta charset="utf-8">\n'
        '  <meta name="viewport" content="width=device-width, initial-scale=1">\n'
        "  <title>Manuscript</title>\n"
        "  <style>\n"
        "    :root { color-scheme: light; }\n"
        "    body { margin: 0; background: #f7f7f4; color: #1f2933; "
        "font: 16px/1.65 Georgia, 'Times New Roman', serif; }\n"
        "    main { max-width: 760px; margin: 0 auto; padding: 48px 28px 72px; "
        "background: #fff; min-height: 100vh; box-sizing: border-box; }\n"
        "    h1, h2, h3 { font-family: Arial, Helvetica, sans-serif; "
        "line-height: 1.25; margin: 2rem 0 0.75rem; color: #102a43; }\n"
        "    h1 { font-size: 2rem; }\n"
        "    h2 { font-size: 1.45rem; border-bottom: 1px solid #d9e2ec; "
        "padding-bottom: 0.35rem; }\n"
        "    h3 { font-size: 1.15rem; }\n"
        "    p { margin: 0 0 1rem; }\n"
        "    a { color: #0b7285; }\n"
        "    blockquote { margin: 1rem 0; padding-left: 1rem; "
        "border-left: 3px solid #9fb3c8; color: #486581; }\n"
        "    code { font-family: Consolas, Monaco, monospace; font-size: 0.95em; }\n"
        "    table { border-collapse: collapse; width: 100%; margin: 0.5rem 0 1rem; "
        "font-size: 0.95em; }\n"
        "    th, td { border: 1px solid #cbd5e0; padding: 0.4rem 0.6rem; "
        "text-align: left; vertical-align: top; }\n"
        "    th { background: #edf2f7; font-weight: 600; }\n"
        "    figure.table-figure, figure.image-figure { margin: 1.25rem 0; }\n"
        "    figcaption.table-caption { margin-bottom: 0.4rem; "
        "font-size: 0.95em; color: #243b53; }\n"
        "    figcaption.image-caption { margin-top: 0.4rem; "
        "font-size: 0.95em; color: #243b53; text-align: center; }\n"
        "    figure.image-figure img { max-width: 100%; display: block; "
        "margin: 0 auto; }\n"
        "    @media print { body { background: #fff; } main { max-width: none; "
        "padding: 0; } a { color: inherit; text-decoration: none; } }\n"
        "  </style>\n"
        "</head>\n"
        "<body>\n"
        f"<main>\n{body}\n</main>\n"
        "</body>\n"
        "</html>\n"
    )
    _write_text(path, html)


def _csl_items(
    shortlist: Sequence[NormalizedSource],
    cited_source_ids: set[str],
    project_language: str = "en",
) -> list[dict[str, object]]:
    """Return CSL-JSON items for the cited sources.

    Each item carries a per-source ``language`` field set to the project
    language so downstream citeproc renderers can pick the right CSL
    style locale (APA for en, GB/T 7714 for zh, JSPS / SIST 02 for ja).
    Source-specific overrides (e.g. an English source cited in a Chinese
    paper) are not detected here — callers wanting per-source locale
    detection should preprocess the shortlist first.
    """
    locale = (project_language or "en").strip().lower() or "en"
    items: list[dict[str, object]] = []
    for source in shortlist:
        if source.source_id not in cited_source_ids:
            continue
        item: dict[str, object] = {
            "id": source.source_id,
            "type": "article-journal",
            "title": source.title,
            "author": [{"literal": author} for author in source.authors],
            "language": locale,
        }
        if source.year is not None:
            item["issued"] = {"date-parts": [[source.year]]}
        if source.venue:
            item["container-title"] = source.venue
        if source.doi:
            item["DOI"] = source.doi
        if source.url:
            item["URL"] = source.url
        items.append(item)
    return items


def _manifest_payload(
    run_dir: Path,
    artifacts: Mapping[str, str],
    project_language: str = "en",
) -> dict[str, object]:
    files: dict[str, dict[str, object]] = {}
    for export_format, relative_path in artifacts.items():
        path = run_dir / relative_path
        files[export_format] = {
            "path": relative_path,
            "sha256": _sha256(path),
            "size_bytes": path.stat().st_size,
        }
    return {
        "files": files,
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "git_sha": _git_sha(),
        "language": (project_language or "en").strip().lower() or "en",
        "citation_style_hint": _citation_style_hint(project_language),
    }


def _citation_style_hint(language: str | None) -> str:
    code = (language or "en").strip().lower()
    if code == "zh":
        # GB/T 7714-2015 numeric is the de-facto Chinese journal standard.
        return "gb-t-7714-2015-numeric"
    if code == "ja":
        # SIST 02 is the dominant style for Japanese academic journals.
        return "sist02"
    return "apa"


def _download_links(run_id: str, manifest: Mapping[str, object]) -> list[dict[str, object]]:
    """PR-379 (2026-05-13): each download link now carries a content-
    hash ``?v=`` query parameter so Cloudflare treats every new
    export as a fresh URL.

    Without the version qualifier, CF kept serving the 4h-cached old
    docx after PR-375/378 deployed even though the on-disk file had
    changed and the API set ``Cache-Control: no-store``. The cached
    cache entry (from the pre-no-store era) persisted under the bare
    URL. Adding ``?v={sha256[:8]}`` makes the URL effectively
    immutable per content: cache hits only when the file actually
    hasn't changed; new content → new URL → CF miss → origin hit
    → ``no-store`` response → never cached.
    """
    files = manifest.get("files")
    if not isinstance(files, dict):
        return []
    links: list[dict[str, object]] = []
    for export_format, payload in files.items():
        if not isinstance(payload, dict):
            continue
        path = payload.get("path")
        if not isinstance(path, str):
            continue
        filename = Path(path).name
        sha = payload.get("sha256")
        version_qualifier = f"?v={str(sha)[:8]}" if isinstance(sha, str) and sha else ""
        links.append(
            {
                "format": export_format,
                "filename": filename,
                "url": f"/api/runs/{run_id}/exports/{filename}{version_qualifier}",
            },
        )
    return links


def _read_synthesizer_source_notes(run_dir: Path) -> dict[str, dict[str, object]]:
    """Load Synthesizer's per-source notes if available.

    Returns ``{source_id: {thesis, method, evidence, limits, ...}}``.
    Empty dict on any I/O / parse error — the literature usage table
    just falls back to the source abstract.
    """
    notes_dir = run_dir / "synthesis" / "source_notes"
    if not notes_dir.exists() or not notes_dir.is_dir():
        return {}
    out: dict[str, dict[str, object]] = {}
    for path in notes_dir.glob("*.json"):
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if not isinstance(payload, dict):
            continue
        source_id = str(payload.get("source_id") or path.stem)
        out[source_id] = payload
    return out


def _resolve_authors(session: Session, project: Project | None) -> list[str]:
    """Return the author byline for this paper.

    Codex-AGREEd fallback chain:
    1. ``project_author`` rows ordered by position (the user's
       explicit author selection for this project)
    2. The user's lazy-bootstrapped ``self``-author, if any
    3. ``User.display_name``
    4. Literal ``"Admin"`` (legacy / test fallback only)
    """
    from sqlalchemy import select as _select

    from autoessay.models import Author, ProjectAuthor

    if project is None:
        return []
    rows = list(
        session.scalars(
            _select(Author)
            .join(ProjectAuthor, ProjectAuthor.author_id == Author.id)
            .where(ProjectAuthor.project_id == project.id)
            .order_by(ProjectAuthor.position.asc()),
        ).all()
    )
    if rows:
        return [a.display_name for a in rows if a.display_name]
    self_author = session.scalar(
        _select(Author)
        .where(Author.user_id == project.user_id, Author.is_self.is_(True))
        .where(Author.deleted_at.is_(None)),
    )
    if self_author is not None:
        return [self_author.display_name]
    user = session.get(User, project.user_id)
    if user is not None:
        name = (user.display_name or "").strip()
        if name:
            return [name]
        return [user.id] if user.id else ["Admin"]
    return ["Admin"]


def _cited_source_ids(claim_map: Sequence[Mapping[str, object]]) -> set[str]:
    cited: set[str] = set()
    for record in claim_map:
        raw_source_ids = record.get("source_ids")
        if not isinstance(raw_source_ids, list):
            continue
        for source_id in raw_source_ids:
            if isinstance(source_id, str) and source_id != "[UNCITED]":
                cited.add(source_id)
    return cited


def _read_sources_json(path: Path) -> list[NormalizedSource]:
    records = _load_json_array(path)
    sources: list[NormalizedSource] = []
    for record in records:
        if not isinstance(record, dict):
            continue
        try:
            sources.append(NormalizedSource.parse_obj(record))
        except ValidationError:
            continue
    return sources


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _git_sha() -> str:
    settings = get_settings()
    try:
        output = subprocess.check_output(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=Path(__file__).resolve().parents[3],
            stderr=subprocess.DEVNULL,
            text=True,
        )
    except (OSError, subprocess.CalledProcessError):
        return settings.git_sha
    return output.strip() or settings.git_sha


def _json_object(value: str) -> dict[str, object]:
    try:
        decoded = json.loads(value)
    except json.JSONDecodeError:
        return {}
    if not isinstance(decoded, dict):
        return {}
    return {key: value for key, value in decoded.items() if isinstance(key, str)}


def _load_json_array(path: Path) -> list[object]:
    if not path.exists():
        return []
    try:
        decoded = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []
    return decoded if isinstance(decoded, list) else []


def _load_json_mapping(path: Path) -> dict[str, object]:
    if not path.exists():
        return {}
    try:
        decoded = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}
    if not isinstance(decoded, dict):
        return {}
    return {key: value for key, value in decoded.items() if isinstance(key, str)}


def _load_jsonl_objects(path: Path) -> list[dict[str, object]]:
    if not path.exists():
        return []
    records: list[dict[str, object]] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            stripped = line.strip()
            if not stripped:
                continue
            try:
                decoded = json.loads(stripped)
            except json.JSONDecodeError:
                continue
            if isinstance(decoded, dict):
                records.append(decoded)
    return records


def _read_optional_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8")


# 2026-05-12 PR-364: manuscript markdown carries LaTeX math (``$$...$$``),
# markdown tables, and ``【待填】`` placeholders introduced by round-0
# stage B's empirical scaffolding. The docx exporter does not render
# math, and the html exporter delegates math to MathJax (browser-side
# only). Add manuscript.tex so authors targeting CNKI / 顶刊 / arXiv
# submission flows get proper LaTeX math + tables and can compile to
# PDF directly. The converter handles the markdown subset the pipeline
# actually emits — heading levels, lists, markdown tables, fenced math
# blocks, inline math, basic emphasis — and escapes LaTeX-special
# characters outside math regions. Anything richer can be hand-edited
# in the produced .tex.

_LATEX_ESCAPE_MAP = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _latex_escape(text: str) -> str:
    """Escape LaTeX-special characters in plain prose. Math regions
    (``$...$`` and ``$$...$$``) and code spans are handled by callers
    BEFORE calling this — once text reaches here it must be treated
    as literal."""
    out: list[str] = []
    for ch in text:
        out.append(_LATEX_ESCAPE_MAP.get(ch, ch))
    return "".join(out)


_INLINE_MATH_RE = re.compile(r"\$([^$\n]+?)\$")


def _latex_escape_prose_preserving_inline_math(text: str) -> str:
    """Escape LaTeX-specials but pass single-dollar inline math
    (``$x$``) through untouched. Display math (``$$...$$``) is handled
    at the block level by ``_write_latex`` and never reaches here.
    Also preserves citation markers ``[N]`` (LaTeX-safe already) and
    ``【待填】`` highlighting (wrapped in colored box at the line level).
    """
    parts: list[str] = []
    cursor = 0
    for match in _INLINE_MATH_RE.finditer(text):
        parts.append(_latex_escape(text[cursor : match.start()]))
        # ``$...$`` is passed through as-is; LaTeX accepts inline math
        # in this exact syntax.
        parts.append(match.group(0))
        cursor = match.end()
    parts.append(_latex_escape(text[cursor:]))
    out = "".join(parts)
    # Highlight 【待填】 placeholders so reviewers see them at a glance.
    out = out.replace("【待填】", r"\textcolor{red}{【待填】}")
    return out


def _emphasis_to_latex(text: str) -> str:
    """Convert markdown emphasis to LaTeX commands. Applied AFTER
    LaTeX escaping (so the backslashes survive)."""
    # **bold**
    text = re.sub(r"\*\*([^*\n]+?)\*\*", r"\\textbf{\1}", text)
    # *italic*
    text = re.sub(r"(?<!\*)\*([^*\n]+?)\*(?!\*)", r"\\textit{\1}", text)
    return text


def _markdown_table_to_latex(rows: list[str], table_number: int | None = None) -> list[str]:
    """Convert a markdown table block (list of ``| ... |`` lines, the
    second of which is the ``|---|---|`` separator) to a LaTeX
    ``tabular`` environment with ``booktabs`` rules.

    PR-375: when ``table_number`` is given, prepend a ``\\caption{表 N}``
    line above the tabular (codex AGREE-WITH-AMENDMENTS amendment 1 —
    table caption goes ABOVE in Chinese journal convention).
    """
    # Strip empty leading/trailing pipes and split into cells per row.
    cells_per_row: list[list[str]] = []
    for row in rows:
        line = row.strip()
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        cells_per_row.append(cells)
    if not cells_per_row:
        return []
    # First row = header, second row = separator (e.g. "|---|---|"); we
    # detect the separator by presence of dashes only.
    header = cells_per_row[0]
    body_start = 1
    if len(cells_per_row) > 1:
        sep_cells = cells_per_row[1]
        if all(re.fullmatch(r":?-+:?", cell) for cell in sep_cells):
            body_start = 2
    num_cols = max(len(header), 1)
    col_spec = "l" * num_cols
    out: list[str] = [
        r"\begin{table}[h]",
        r"\centering",
    ]
    if table_number is not None:
        # Caption ABOVE the tabular (Chinese journal convention,
        # codex amendment 1). ``\caption*`` would skip auto-numbering;
        # using a literal ``表 N`` plus ``\captionsetup{labelformat=
        # empty}`` is overkill — just inline the literal text in a
        # bold paragraph above the centering block instead. We pick
        # the explicit-literal route so the .tex stays self-contained
        # without an extra package.
        out.append(r"\noindent\textbf{表 " + str(table_number) + r"}\\[2pt]")
    out.extend(
        [
            r"\begin{tabular}{" + col_spec + r"}",
            r"\toprule",
            " & ".join(_latex_escape_prose_preserving_inline_math(c) for c in header) + r" \\",
            r"\midrule",
        ],
    )
    for body_row in cells_per_row[body_start:]:
        # Pad short rows to match header width.
        padded = body_row + [""] * (num_cols - len(body_row))
        out.append(
            " & ".join(_latex_escape_prose_preserving_inline_math(c) for c in padded[:num_cols])
            + r" \\"
        )
    out.extend(
        [
            r"\bottomrule",
            r"\end{tabular}",
            r"\end{table}",
            "",
        ]
    )
    return out


def _latex_preamble(language: str) -> list[str]:
    """Build the LaTeX preamble. For zh/ja we use ``ctexart`` so CJK
    characters render without manual font configuration; for other
    languages we use the plain ``article`` class with ``inputenc``
    UTF-8 and a few standard packages."""
    is_cjk = (language or "").strip().lower() in {"zh", "ja", "zh-cn", "zh-tw"}
    doc_class = (
        r"\documentclass[12pt,a4paper]{ctexart}"
        if is_cjk
        else r"\documentclass[12pt,a4paper]{article}"
    )
    common = [
        doc_class,
        r"\usepackage{amsmath}",
        r"\usepackage{amssymb}",
        r"\usepackage{booktabs}",
        r"\usepackage{graphicx}",
        r"\usepackage{xcolor}",
        r"\usepackage{geometry}",
        r"\geometry{margin=2.5cm}",
        r"\usepackage{hyperref}",
        r"\hypersetup{colorlinks=true, linkcolor=blue, citecolor=blue, urlcolor=blue}",
    ]
    if not is_cjk:
        common.insert(1, r"\usepackage[utf8]{inputenc}")
    return common


def _write_latex(path: Path, manuscript: str, language: str = "en") -> None:
    """Convert manuscript markdown to a standalone .tex source. The
    output document is compileable with ``xelatex`` (ctexart for zh)
    and contains: title from the first ``# `` heading, then
    ``\\section / \\subsection`` for ``## / ###`` headings, ``itemize``
    for ``- ...`` bullets, ``tabular`` for markdown tables, ``\\[ ...
    \\]`` for ``$$...$$`` display math (single-dollar inline math passes
    through unchanged). LaTeX-special characters in prose are escaped;
    ``【待填】`` placeholders are highlighted in red so reviewers can
    spot them at a glance.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    out: list[str] = list(_latex_preamble(language))

    title: str = ""
    body: list[str] = []
    lines = manuscript.splitlines()
    i = 0
    in_display_math = False
    math_buffer: list[str] = []
    list_buffer: list[str] = []
    # PR-375: auto-numbered captions across the document. Tables get a
    # caption ABOVE (Chinese convention), figures get one BELOW (codex
    # amendment 1).
    table_counter = 0
    figure_counter = 0

    def flush_list() -> None:
        if not list_buffer:
            return
        body.append(r"\begin{itemize}")
        for li in list_buffer:
            body.append(
                r"  \item " + _emphasis_to_latex(_latex_escape_prose_preserving_inline_math(li))
            )
        body.append(r"\end{itemize}")
        body.append("")
        list_buffer.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Display-math block ``$$...$$`` (may span multiple lines).
        if in_display_math:
            if "$$" in line:
                # Pull text before the closing ``$$`` into the math
                # buffer, then close.
                end = line.index("$$")
                math_buffer.append(line[:end])
                body.append(
                    r"\[" + " ".join(part.strip() for part in math_buffer if part.strip()) + r"\]"
                )
                body.append("")
                math_buffer.clear()
                in_display_math = False
                rest = line[end + 2 :].strip()
                if rest:
                    # Anything after the closing ``$$`` on the same line
                    # is treated as a fresh prose line on the next loop.
                    lines.insert(i + 1, rest)
                i += 1
                continue
            math_buffer.append(line)
            i += 1
            continue

        if "$$" in stripped:
            flush_list()
            # Opening ``$$`` — may also close on the same line.
            start = stripped.index("$$")
            after = stripped[start + 2 :]
            if "$$" in after:
                end = after.index("$$")
                math_content = after[:end]
                body.append(r"\[" + math_content.strip() + r"\]")
                body.append("")
                tail = after[end + 2 :].strip()
                if tail:
                    lines.insert(i + 1, tail)
                i += 1
                continue
            # Multi-line open
            if after.strip():
                math_buffer.append(after.strip())
            in_display_math = True
            i += 1
            continue

        # Skip raw HTML anchors injected by the pipeline (e.g.
        # ``<a id="s01_p01"></a>``); LaTeX has its own \label mechanism
        # but we keep this simple by stripping them.
        if re.fullmatch(r'<a id="[^"]+"></a>', stripped):
            i += 1
            continue

        if stripped.startswith("- "):
            list_buffer.append(stripped[2:].strip())
            i += 1
            continue
        else:
            flush_list()

        if _is_table_row(stripped):
            # Collect contiguous table rows (PR-381: ``_is_table_row``
            # tolerates trailing ``[N]`` citation markers).
            table_rows: list[str] = []
            while i < len(lines):
                sub = lines[i].strip()
                if _is_table_row(sub):
                    table_rows.append(_strip_trailing_table_citations(sub))
                    i += 1
                    continue
                # Allow the second-line separator through even though
                # it's not a "table row" by our predicate.
                if _MD_TABLE_SEPARATOR_RE.match(sub):
                    table_rows.append(sub)
                    i += 1
                    continue
                break
            table_counter += 1
            body.extend(_markdown_table_to_latex(table_rows, table_number=table_counter))
            continue

        # PR-375: ``![alt](url)`` markdown image → ``\includegraphics``
        # placeholder + ``图 N`` caption BELOW (codex amendment 1).
        # Codex amendment 6: don't fetch remote URLs; just emit the
        # caption + a placeholder so the .tex still compiles when the
        # image isn't on disk.
        img_match = _MD_IMAGE_RE.fullmatch(stripped)
        if img_match:
            figure_counter += 1
            alt = img_match.group(1)
            url = img_match.group(2)
            body.append(r"\begin{figure}[h]")
            body.append(r"\centering")
            # \includegraphics will fail to compile if the path
            # doesn't exist, so wrap in IfFileExists to degrade
            # gracefully. The bracket form keeps the .tex compileable
            # even without the asset (placeholder text instead).
            url_escaped = _latex_escape(url)
            body.append(
                r"\IfFileExists{"
                + url_escaped
                + r"}{\includegraphics[width=\linewidth]{"
                + url_escaped
                + r"}}{\fbox{\parbox{0.8\linewidth}{\centering 图 "
                + str(figure_counter)
                + r" 占位 (URL: "
                + url_escaped
                + r")}}}",
            )
            cap_text = (
                _emphasis_to_latex(_latex_escape_prose_preserving_inline_math(alt)) if alt else ""
            )
            body.append(
                r"\\[2pt]\textbf{图 "
                + str(figure_counter)
                + r"}"
                + ("  " + cap_text if cap_text else "")
            )
            body.append(r"\end{figure}")
            body.append("")
            i += 1
            continue

        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if not title:
                title = _latex_escape_prose_preserving_inline_math(heading)
            else:
                body.append(
                    r"\section*{"
                    + _emphasis_to_latex(_latex_escape_prose_preserving_inline_math(heading))
                    + r"}"
                )
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            body.append(
                r"\section{"
                + _emphasis_to_latex(_latex_escape_prose_preserving_inline_math(heading))
                + r"}"
            )
            i += 1
            continue
        if stripped.startswith("### "):
            heading = stripped[4:].strip()
            body.append(
                r"\subsection{"
                + _emphasis_to_latex(_latex_escape_prose_preserving_inline_math(heading))
                + r"}"
            )
            i += 1
            continue
        if stripped.startswith("#### "):
            heading = stripped[5:].strip()
            body.append(
                r"\subsubsection{"
                + _emphasis_to_latex(_latex_escape_prose_preserving_inline_math(heading))
                + r"}"
            )
            i += 1
            continue

        # Empty line — separate paragraphs.
        if not stripped:
            body.append("")
            i += 1
            continue

        # Default: prose paragraph line.
        body.append(_emphasis_to_latex(_latex_escape_prose_preserving_inline_math(stripped)))
        i += 1

    flush_list()
    # If the document ended with an unterminated ``$$`` block, flush
    # what we have so the output is at least syntactically closeable.
    if in_display_math and math_buffer:
        body.append(r"\[" + " ".join(math_buffer) + r"\]")

    if title:
        out.extend([r"\title{" + title + r"}", r"\author{}", r"\date{}"])
    out.append(r"\begin{document}")
    if title:
        out.append(r"\maketitle")
    out.extend(body)
    out.append(r"\end{document}")
    _write_text(path, "\n".join(out) + "\n")


def _write_json(path: Path, payload: object) -> None:
    _write_text(
        path,
        json.dumps(payload, indent=2, sort_keys=True, ensure_ascii=False) + "\n",
    )


def _write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(value, encoding="utf-8")
    temporary.replace(path)
