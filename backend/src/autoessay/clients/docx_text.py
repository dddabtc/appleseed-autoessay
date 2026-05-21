"""DOCX text extraction for local prior-paper ingestion."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from autoessay.clients.pdf_text import MIN_EXTRACTED_CHARS, PoorExtraction


def extract_text(docx_bytes: bytes, source_id: str | None = None) -> str:
    label = source_id or "unknown"
    try:
        document = Document(BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001 - caller needs a deterministic warning class.
        raise PoorExtraction(f"DOCX text extraction failed for {label}: {exc}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    extracted = "\n\n".join(parts)
    if len(extracted.strip()) < MIN_EXTRACTED_CHARS:
        raise PoorExtraction(
            f"DOCX text extraction yielded fewer than {MIN_EXTRACTED_CHARS} characters",
        )
    return extracted
